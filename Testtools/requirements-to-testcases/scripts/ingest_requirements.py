#!/usr/bin/env python3
"""Ingest multi-format requirement documents and emit normalized requirements.json.

Supported inputs:
- .md, .txt
- .docx (python-docx)
- .xlsx/.xlsm/.xltx/.xltm (openpyxl, including embedded images OCR)
- .jpg/.jpeg/.png (OCR via pytesseract+Pillow, fallback easyocr)
"""

from __future__ import annotations

import argparse
import io
import json
import tempfile
from pathlib import Path
from typing import Iterable, List

from extract_requirements import extract_requirements

TEXT_EXTS = {".md", ".txt"}
EXCEL_EXTS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png"}
DOCX_EXTS = {".docx"}
SUPPORTED_EXTS = TEXT_EXTS | EXCEL_EXTS | IMAGE_EXTS | DOCX_EXTS


def _ensure_tesseract_cmd(pytesseract_module) -> None:
    current = getattr(pytesseract_module.pytesseract, "tesseract_cmd", "") or ""
    if current and Path(current).exists():
        return

    candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            pytesseract_module.pytesseract.tesseract_cmd = candidate
            return


def _resolve_tesseract_lang_and_config(ocr_lang: str) -> tuple[str, str]:
    user_tessdata = Path.home() / "tessdata"
    if "chi_sim" not in ocr_lang or not user_tessdata.exists():
        return ocr_lang, ""

    available = {p.stem for p in user_tessdata.glob("*.traineddata")}
    requested = [x.strip() for x in ocr_lang.split("+") if x.strip()]

    # Prefer using all requested languages when model files are available.
    if requested and all(lang in available for lang in requested):
        return ocr_lang, f"--tessdata-dir {user_tessdata}"

    # Fallback: ensure Chinese OCR still works when only chi_sim is present.
    if "chi_sim" in available:
        return "chi_sim", f"--tessdata-dir {user_tessdata}"

    return ocr_lang, ""


def _easyocr_langs(ocr_lang: str) -> List[str]:
    langs: List[str] = []
    if "chi_sim" in ocr_lang:
        langs.append("ch_sim")
    if "eng" in ocr_lang:
        langs.append("en")
    if not langs:
        langs = ["en"]
    return langs


def _ocr_image_bytes(image_bytes: bytes, ocr_lang: str) -> str:
    # Backend 1: pytesseract + Pillow
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore

        _ensure_tesseract_cmd(pytesseract)
        tess_lang, tess_cfg = _resolve_tesseract_lang_and_config(ocr_lang)
        with Image.open(io.BytesIO(image_bytes)) as img:
            return pytesseract.image_to_string(img, lang=tess_lang, config=tess_cfg)
    except Exception:
        pass

    # Backend 2: easyocr
    try:
        import easyocr  # type: ignore

        langs = _easyocr_langs(ocr_lang)
        reader = easyocr.Reader(langs, gpu=False)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=True) as tmp:
            tmp.write(image_bytes)
            tmp.flush()
            result = reader.readtext(tmp.name, detail=0, paragraph=True)
        return "\n".join(result)
    except Exception:
        pass

    # Backend 3: rapidocr-onnxruntime
    try:
        from rapidocr_onnxruntime import RapidOCR  # type: ignore

        engine = RapidOCR()
        result, _ = engine(image_bytes)
        if not result:
            return ""
        return "\n".join([item[1] for item in result if len(item) >= 2 and item[1]])
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "No OCR backend available. Install one of: "
            "(1) pip install pillow pytesseract and install Tesseract OCR binary, "
            "(2) pip install easyocr, or "
            "(3) pip install rapidocr-onnxruntime"
        ) from exc


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_docx_file(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency for .docx parsing: python-docx. Install with: pip install python-docx"
        ) from exc

    doc = Document(str(path))
    lines: List[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def read_excel_file(path: Path, ocr_lang: str) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore
        from openpyxl.utils.cell import get_column_letter  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency for .xlsx parsing: openpyxl. Install with: pip install openpyxl"
        ) from exc

    # read_only=False is required because embedded images are not exposed in read_only mode.
    wb = load_workbook(filename=str(path), read_only=False, data_only=True)
    lines: List[str] = []

    for ws in wb.worksheets:
        lines.append(f"[Sheet] {ws.title}")

        # Cell text
        for row in ws.iter_rows(values_only=True):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                lines.append(" | ".join(values))

        # Embedded images OCR
        images = list(getattr(ws, "_images", []) or [])
        for idx, img in enumerate(images, start=1):
            anchor_cell = "unknown"
            anchor = getattr(img, "anchor", None)
            if anchor is not None and hasattr(anchor, "_from"):
                row = getattr(anchor._from, "row", 0) + 1
                col = getattr(anchor._from, "col", 0) + 1
                anchor_cell = f"{get_column_letter(col)}{row}"

            image_bytes = b""
            try:
                if hasattr(img, "_data"):
                    image_bytes = img._data()
                else:
                    ref = getattr(img, "ref", None)
                    if isinstance(ref, (str, Path)):
                        image_bytes = Path(ref).read_bytes()
                    elif hasattr(ref, "read"):
                        image_bytes = ref.read()
            except Exception:
                image_bytes = b""

            if not image_bytes:
                lines.append(f"[Sheet Image] {ws.title}!{anchor_cell} # {idx}: <image extracted failed>")
                continue

            try:
                ocr_text = _ocr_image_bytes(image_bytes, ocr_lang=ocr_lang).strip()
            except Exception as exc:
                lines.append(f"[Sheet Image] {ws.title}!{anchor_cell} # {idx}: <ocr unavailable: {exc}>")
                continue

            if ocr_text:
                lines.append(f"[Sheet Image] {ws.title}!{anchor_cell} # {idx}: {ocr_text}")
            else:
                lines.append(f"[Sheet Image] {ws.title}!{anchor_cell} # {idx}: <ocr empty>")

    return "\n".join(lines)


def read_image_file(path: Path, ocr_lang: str) -> str:
    return _ocr_image_bytes(path.read_bytes(), ocr_lang=ocr_lang)


def extract_text(path: Path, ocr_lang: str) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTS:
        return read_text_file(path)
    if suffix in DOCX_EXTS:
        return read_docx_file(path)
    if suffix in EXCEL_EXTS:
        return read_excel_file(path, ocr_lang=ocr_lang)
    if suffix in IMAGE_EXTS:
        return read_image_file(path, ocr_lang=ocr_lang)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def collect_files(input_path: Path) -> List[Path]:
    if input_path.is_file():
        if input_path.suffix.lower() not in SUPPORTED_EXTS:
            raise ValueError(f"Unsupported file type: {input_path.suffix}")
        return [input_path]

    if input_path.is_dir():
        files = sorted(
            [p for p in input_path.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS]
        )
        if not files:
            raise ValueError(f"No supported files found under: {input_path}")
        return files

    raise FileNotFoundError(f"Input path not found: {input_path}")


def ingest(files: Iterable[Path], ocr_lang: str) -> List[dict]:
    all_reqs: List[dict] = []
    counter = 1

    for f in files:
        raw = extract_text(f, ocr_lang=ocr_lang)
        reqs = extract_requirements(raw)

        for req in reqs:
            req["requirement_id"] = f"REQ-{counter:03d}"
            req["source_file"] = str(f)
            all_reqs.append(req)
            counter += 1

    return all_reqs


def main() -> None:
    ap = argparse.ArgumentParser(
        description="Parse .docx/.xlsx/.jpg/.png/.md/.txt requirements and generate requirements.json"
    )
    ap.add_argument("input", help="Input file or directory")
    ap.add_argument("-o", "--output", default="requirements.json", help="Output JSON path")
    ap.add_argument(
        "--ocr-lang",
        default="chi_sim+eng",
        help="OCR language for pytesseract (default: chi_sim+eng)",
    )
    ap.add_argument("--raw-text-out", default="", help="Optional path to save merged extracted text")
    args = ap.parse_args()

    input_path = Path(args.input)
    files = collect_files(input_path)

    all_reqs = ingest(files, ocr_lang=args.ocr_lang)

    output = {
        "requirements": all_reqs,
        "metadata": {
            "input": str(input_path),
            "source_files": [str(f) for f in files],
            "count": len(all_reqs),
        },
    }

    Path(args.output).write_text(json.dumps(output, indent=2, ensure_ascii=False), encoding="utf-8")

    if args.raw_text_out:
        merged = []
        for f in files:
            merged.append(f"### SOURCE: {f}")
            merged.append(extract_text(f, ocr_lang=args.ocr_lang))
            merged.append("")
        Path(args.raw_text_out).write_text("\n".join(merged), encoding="utf-8")

    print(f"Parsed {len(files)} file(s), extracted {len(all_reqs)} requirement(s) -> {args.output}")


if __name__ == "__main__":
    main()
