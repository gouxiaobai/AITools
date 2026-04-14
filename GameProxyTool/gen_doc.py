"""
生成 GameProxyTool 使用手册.docx
依赖：pip install python-docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 全局字体 ────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

def set_heading(paragraph, level):
    """设置标题中文字体"""
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

def heading(text, level):
    p = doc.add_heading(text, level=level)
    set_heading(p, level)
    return p

def para(text="", bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "微软雅黑"
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p

def code_block(text):
    """等宽字体代码块（带底色模拟）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    # 灰色底色通过段落着色
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)
    for line in text.strip().split("\n"):
        if p.runs:
            p.add_run("\n" + line).font.name = "Courier New"
        else:
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # 表头
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "微软雅黑"
        cell.paragraphs[0].runs[0].font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        # 表头底色
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E1F2")
        tc_pr.append(shd)
    # 数据行
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = "微软雅黑"
                run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    # 列宽
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # 表后空行
    return t

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("GameProxyTool 使用手册")
run.bold = True
run.font.size = Pt(22)
run.font.name = "微软雅黑"
run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ["适用项目：BlackMist（Unity + HybridCLR + xLua）",
             "工具版本：基于 Python 3.10+ / PyInstaller 打包",
             "文档日期：2026-03"]:
    run = sub_p.add_run(line + "\n")
    run.font.size = Pt(10)
    run.font.name = "微软雅黑"
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 一、背景与目的
# ════════════════════════════════════════════════════════════════════════════
heading("一、背景与目的", 1)
para("BlackMist 游戏使用自定义二进制 TCP 协议与服务器通信，常规抓包工具（Charles、Fiddler）无法解析数据包内容，也无法对接口进行重放测试。")
para("本工具针对该项目协议格式，实现：")
bullet("透明拦截游戏客户端与服务器之间的全部 TCP 数据包")
bullet("实时解码并以可读 JSON 格式展示")
bullet("支持对任意接口进行重放、修改后重发、自定义发包")
hr()

# ════════════════════════════════════════════════════════════════════════════
# 二、功能范围
# ════════════════════════════════════════════════════════════════════════════
heading("二、功能范围", 1)
heading("2.1 支持的功能", 2)
add_table(
    ["功能", "说明"],
    [
        ["实时抓包",      "捕获客户端↔服务器全部 TCP 数据包，实时展示"],
        ["协议解码",      "自动解密（XOR）并反序列化（MessagePack），显示为 JSON"],
        ["数据包重放",    "选中历史包，一键重发到服务器"],
        ["修改后重放",    "编辑 JSON 字段后再发送，用于测试边界条件"],
        ["自定义发包",    "填写 Action 号和参数，直接向服务器发送任意请求"],
        ["快速模板",      "内置月卡、周卡、邮件等常用接口模板，一键填充"],
        ["SSO 自动拦截",  "自动修改 SSO 响应中的游戏服务器地址，无需改任何服务端配置"],
        ["过滤与搜索",    "按 Action 号、收发方向、内容关键词过滤数据包"],
        ["数据导出",      "将当前过滤结果导出为 JSON 文件"],
        ["数据包备注",    "对关键包打标记备注"],
        ["Token 管理",   "自动从流量中提取 Token；支持手动设置"],
    ],
    col_widths=[4, 12]
)

heading("2.2 不支持的功能", 2)
add_table(
    ["限制", "原因"],
    [
        ["HTTPS SSO 环境",         "工具 HTTP 拦截仅支持明文 HTTP；Release 包走 HTTPS SSO，需额外处理证书"],
        ["SSO 响应已加密时的自动解析", "若服务端对 SSO 响应做了 AES 加密，自动修改 s_url_list 会失败"],
        ["iOS 真机",               "未测试；理论上设置 WiFi 代理后可用，但需确认"],
        ["地图服务器流量",           "地图使用独立 TCP 连接（Action 10042/10043），当前工具仅代理主连接"],
    ],
    col_widths=[5, 11]
)
hr()

# ════════════════════════════════════════════════════════════════════════════
# 三、协议分析
# ════════════════════════════════════════════════════════════════════════════
heading("三、协议分析", 1)
heading("3.1 通信架构", 2)
code_block(
    "游戏客户端\n"
    "  │\n"
    "  ├─ HTTP ──→ SSO 服务器（ssoGetServerList）\n"
    "  │           └─ 返回游戏 TCP 服务器地址列表（s_url_list）\n"
    "  │\n"
    "  └─ TCP ───→ 游戏服务器（自定义二进制协议）"
)

heading("3.2 TCP 数据包格式", 2)
para("客户端发包（C2S）：", bold=True)
code_block("[4字节 包长度，大端序] + [XOR加密的 MessagePack 数据]")
para("服务端收包（S2C）：", bold=True)
code_block("[4字节 包长度，大端序] + [1字节 压缩标志] + [XOR加密的 MessagePack 数据]")
para("注：压缩标志 = 1 时，数据部分为 GZip 压缩后再 XOR 加密。", italic=True)

heading("3.3 加密方式", 2)
code_block(
    "XOR 密钥 = max(1, 数据长度 % 256)\n"
    "对数据的每个字节执行: byte ^ key\n"
    "（加密与解密使用同一算法，对称）"
)

heading("3.4 消息结构（解码后）", 2)
code_block(
    '{\n'
    '  "action": 19660,\n'
    '  "commReq": {\n'
    '    "sid": 42,\n'
    '    "token": "xxxxxxxx",\n'
    '    "cv": 1\n'
    '  },\n'
    '  "其他业务字段": "..."\n'
    '}'
)
bullet("action：接口协议号，唯一标识接口类型")
bullet("commReq.sid：消息序号（客户端自增）")
bullet("commReq.token：用户认证 Token，从登录响应中获取")

heading("3.5 SSO 地址配置（各环境）", 2)
add_table(
    ["环境", "SSO 地址", "协议"],
    [
        ["Debug",           "http://47.95.196.154:10801/api2/",            "HTTP ✅"],
        ["OnlineTest",      "http://aoo_dev_sso.ageofzorigins.cn:10803/api2/", "HTTP ✅"],
        ["ATestOnlineTest", "http://47.95.196.154:10802/api2/",            "HTTP ✅"],
        ["Release",         "https://sso-elb.camelgames-aoz.com/",        "HTTPS ❌"],
    ],
    col_widths=[3.5, 9, 3.5]
)
hr()

# ════════════════════════════════════════════════════════════════════════════
# 四、架构设计
# ════════════════════════════════════════════════════════════════════════════
heading("四、架构设计", 1)
heading("4.1 工具组成", 2)
code_block(
    "GameProxyTool/\n"
    "├── main.py             # 入口，参数解析，组件启动\n"
    "├── protocol.py         # 编解码：XOR + MessagePack + GZip\n"
    "├── store.py            # 数据包存储，Token 提取\n"
    "├── proxy.py            # TCP MITM 代理核心\n"
    "├── http_interceptor.py # HTTP 代理，自动修改 SSO 响应\n"
    "├── web.py              # Flask Web API + SSE 实时推送\n"
    "├── static/\n"
    "│   └── index.html      # Web 控制台（单页应用）\n"
    "└── requirements.txt"
)

heading("4.2 工作流程", 2)
code_block(
    "手机\n"
    " │\n"
    " │ (1) WiFi HTTP 代理 → PC:8080\n"
    " │\n"
    " ├─ SSO HTTP 请求 ──→ [HTTP拦截器 :8080]\n"
    " │                         │\n"
    " │                         ├─ 转发到真实 SSO 服务器\n"
    " │                         ├─ 读取响应中的 s_url_list（真实游戏服务器地址）\n"
    " │                         ├─ 将 s_url_list 替换为 PC:18080\n"
    " │                         └─ 通知 TCP 代理设置转发目标\n"
    " │\n"
    " └─ 游戏 TCP 连接 ──→ [TCP 代理 :18080]\n"
    "                           │\n"
    "                           ├─ 解码 C2S 包 → 存入 PacketStore\n"
    "                           ├─ 转发到真实游戏服务器\n"
    "                           ├─ 解码 S2C 包 → 存入 PacketStore\n"
    "                           └─ 转发回手机\n"
    "\n"
    "浏览器 ──→ [Web 控制台 :8888]\n"
    "              └─ SSE 实时展示 / REST API 发包"
)
hr()

# ════════════════════════════════════════════════════════════════════════════
# 五、操作手册
# ════════════════════════════════════════════════════════════════════════════
heading("五、操作手册", 1)
heading("5.1 环境要求", 2)
bullet("PC 系统：Windows 10/11")
bullet("手机：Android（iOS 理论可用，未测试）")
bullet("网络：手机与 PC 连接同一 WiFi")

heading("5.2 启动工具", 2)
para("双击运行 GameProxyTool.exe，控制台显示：")
code_block(
    "============================================================\n"
    "  GameProxyTool - TCP 协议调试工具\n"
    "============================================================\n"
    "  游戏 TCP 服务器 IP（直接回车 = 自动从 SSO 获取）\n"
    "  请输入："
)
add_table(
    ["情况", "操作"],
    [
        ["已知游戏服务器 IP",    "输入 IP 或 IP:端口，回车"],
        ["不知道服务器 IP（常用）", "直接回车，工具自动从 SSO 响应中提取"],
    ],
    col_widths=[5, 11]
)
para("启动成功后控制台显示：")
code_block(
    "  ★ 手机 WiFi 代理 → 192.168.x.x:8080\n"
    "  ★ 设好后重启游戏即可开始抓包"
)

heading("5.3 配置手机 WiFi 代理", 2)
bullet("手机进入 设置 → WiFi → 长按当前网络 → 修改网络")
bullet("显示高级选项，代理选择 手动")
bullet("主机名：控制台显示的 PC 局域网 IP（如 192.168.1.100）")
bullet("端口：8080")
bullet("保存")
p = doc.add_paragraph()
run = p.add_run("\u26a0\ufe0f  测试完成后记得将代理改回\u201c无\u201d，否则手机断网。")
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)
run.font.name = "微软雅黑"
run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

heading("5.4 打开 Web 控制台", 2)
para("在 PC 浏览器中打开：")
code_block("http://localhost:8888")
para("未建立连接前，页面顶部显示配置引导。")

heading("5.5 开始抓包", 2)
para("重启游戏客户端。游戏启动时：")
bullet("(1) 发出 SSO HTTP 请求 → 工具拦截，替换游戏服务器地址")
bullet("(2) 游戏 TCP 连接打到工具代理端口")
bullet("(3) Web 控制台实时显示数据包流")
para("控制台日志示例：")
code_block(
    "[HTTP拦截] s_url_list ['10.0.0.5:8080'] → [192.168.1.100:18080]\n"
    "[Proxy] 游戏服务器已发现：10.0.0.5:8080\n"
    "[Proxy] 新连接：('192.168.1.x', 54321)"
)

heading("5.6 Web 控制台功能说明", 2)
heading("数据包列表（左侧）", 3)
add_table(
    ["列", "说明"],
    [
        ["#",            "包序号"],
        ["C2S / S2C",    "绿色=客户端发出，红色=服务端返回"],
        ["Action 号",    "接口协议号"],
        ["接口名称",      "自动映射的中文/英文名称"],
        ["时间",          "抓包时间"],
    ],
    col_widths=[4, 12]
)
para("过滤操作：")
bullet("顶部输入 Action 号 → 只显示该接口的包")
bullet("选择方向 → 只看 C2S 或 S2C")
bullet("搜索框 → 按包内容关键词过滤")

heading("查看 Tab", 3)
para("点击左侧数据包，右侧展示语法高亮的 JSON 内容。")

heading("编辑/重放 Tab", 3)
bullet("点击左侧某个 C2S 包 → 点击 编辑重放")
bullet("修改 JSON 字段值")
bullet("点击 重放（发送到服务器）")
bullet("观察服务端响应")
para("注：commReq（sid/token）会自动重新生成，无需手动处理。", italic=True)

heading("发包 Tab", 3)
para("快速模板：点击预设的接口按钮，自动填充 Action 和参数。")
para("自定义发包：")
bullet("填写 Action 号")
bullet("填写参数 JSON（不需要填 commReq，自动添加）")
bullet("点击 发送")
hr()

# ════════════════════════════════════════════════════════════════════════════
# 六、已知接口清单
# ════════════════════════════════════════════════════════════════════════════
heading("六、已知接口清单", 1)
heading("月卡", 2)
add_table(
    ["Action", "接口名", "参数", "说明"],
    [
        ["19660", "MonthCardReceiveDailyReward",  "无",                     "领取月卡每日奖励"],
        ["19661", "MonthAndWeekCardAllReward",     '{"activityId": 0}',      "一键领取月卡+周卡所有奖励"],
    ],
    col_widths=[2, 6, 4, 4]
)

heading("周卡", 2)
add_table(
    ["Action", "接口名", "参数", "说明"],
    [
        ["19650", "WeekCardSyncInfo",           '{"activityId": 0}',               "同步周卡信息"],
        ["19651", "WeekCardReceiveFreeBox",      '{"activityId": 0}',               "领取周卡免费宝箱"],
        ["19652", "WeekCardReceiveDailyReward",  '{"activityId": 0, "cardId": 0}',  "领取单个周卡每日奖励"],
        ["19653", "WeekCardSyncItems",           '{"activityId": 0}',               "同步周卡物品"],
        ["19654", "WeekCardReceiveAllReward",    '{"activityId": 0}',               "一键领取所有周卡奖励"],
    ],
    col_widths=[2, 5.5, 4.5, 4]
)

heading("邮件", 2)
add_table(
    ["Action", "接口名", "参数", "说明"],
    [
        ["50001", "RewardMail",    '{"mailIdList": [0]}',                      "领取邮件附件奖励（支持批量）"],
        ["50002", "DeleteMails",   '{"mails": [0]}',                           "删除邮件"],
        ["50005", "CollectMail",   '{"mailId": 0}',                            "收藏/取消收藏邮件"],
        ["50006", "SyncMailInfo",  '{"maxMailId": 0, "mailSyncVersion": 0}',   "同步邮件信息"],
    ],
    col_widths=[2, 4, 6, 4]
)

heading("通用", 2)
add_table(
    ["Action", "接口名", "说明"],
    [
        ["0",     "VoidResponse", "空响应"],
        ["1",     "Sync",         "数据同步"],
        ["3000",  "ChatSend",     "发送聊天"],
        ["10042", "MapTcpLogin",  "地图服务器登录"],
        ["20000", "SyncMapData",  "同步地图数据"],
    ],
    col_widths=[2, 5, 9]
)
hr()

# ════════════════════════════════════════════════════════════════════════════
# 七、典型测试场景
# ════════════════════════════════════════════════════════════════════════════
heading("七、典型测试场景", 1)

heading("场景 1：验证月卡每日奖励防重复", 2)
bullet("正常领取一次月卡奖励（Action 19660），观察 S2C 响应是否成功")
bullet("在控制台找到该 C2S 包，点击 重放")
bullet("观察 S2C 响应：是否返回错误码，防止重复领取")

heading("场景 2：邮件越权领取测试", 2)
bullet("在 发包 Tab 选择《邮件领取奖励》模板")
bullet("将 mailIdList 改为其他账号的邮件 ID")
bullet("发送，观察服务端是否返回权限错误")

heading("场景 3：周卡已过期仍领取", 2)
bullet("GM 后台设置周卡过期时间为过去")
bullet("在发包 Tab 发送 Action 19652，填入对应 activityId 和 cardId")
bullet("验证服务端是否拒绝请求")

heading("场景 4：批量邮件奖励压测", 2)
bullet("选择邮件领取模板")
bullet("mailIdList 填入大量 mailId，如 [1,2,3,...,100]")
bullet("观察服务端是否有数量限制或性能异常")
hr()

# ════════════════════════════════════════════════════════════════════════════
# 八、常见问题
# ════════════════════════════════════════════════════════════════════════════
heading("八、常见问题", 1)
add_table(
    ["问题", "原因", "解决"],
    [
        ["手机设代理后无法上网",           "工具未运行或端口被占用",         "确认 GameProxyTool.exe 已运行且 8080 端口未被占用"],
        ["Web 控制台无法打开",            "端口 8888 被占用",              "用 --web-port 9999 换端口"],
        ["无法抓到包",                   "SSO 走 HTTPS（Release 包）",    "切换为 Debug 或 OnlineTest 包"],
        ["数据包解析失败 _parse_error",   "SSO 修改失败，游戏连到了真实服务器", "检查手机 WiFi 代理是否设置正确，重启游戏"],
        ["控制台一直等待游戏服务器地址",    "SSO 请求未经过工具代理",          "确认手机 WiFi 代理主机名填的是 PC 局域网 IP（非 127.0.0.1）"],
    ],
    col_widths=[4, 4.5, 7.5]
)
hr()

# ════════════════════════════════════════════════════════════════════════════
# 九、命令行参数（高级用法）
# ════════════════════════════════════════════════════════════════════════════
heading("九、命令行参数（高级用法）", 1)
add_table(
    ["参数", "说明", "默认值"],
    [
        ["--server-host", "游戏 TCP 服务器 IP（可选，不填则自动从 SSO 获取）", "自动获取"],
        ["--server-port", "游戏 TCP 服务器端口",                             "8080"],
        ["--listen-port", "本地 TCP 代理端口",                              "18080"],
        ["--http-port",   "HTTP 拦截代理端口",                              "8080"],
        ["--web-port",    "Web 控制台端口",                                "8888"],
    ],
    col_widths=[4, 9, 3]
)
para("多环境并行测试示例：")
code_block(
    "# 测试服 A\n"
    "GameProxyTool.exe --server-host 10.0.0.1 --listen-port 18080 --http-port 8080 --web-port 8888\n"
    "\n"
    "# 测试服 B（另开一个窗口）\n"
    "GameProxyTool.exe --server-host 10.0.0.2 --listen-port 18081 --http-port 8081 --web-port 8889"
)
hr()

# ════════════════════════════════════════════════════════════════════════════
# 十、源码重新打包
# ════════════════════════════════════════════════════════════════════════════
heading("十、源码重新打包", 1)
para("如需修改工具后重新打包：")
code_block(
    "cd Tools/GameProxyTool\n"
    "# 安装依赖（首次）\n"
    "pip install flask msgpack pyinstaller\n"
    "\n"
    "# 打包（双击 build.bat 或手动执行）\n"
    "pyinstaller --clean --onefile --name \"GameProxyTool\" \\\n"
    "  --add-data \"static;static\" --collect-all flask main.py\n"
    "\n"
    "# 输出文件\n"
    "dist/GameProxyTool.exe"
)

# ════════════════════════════════════════════════════════════════════════════
# 保存
# ════════════════════════════════════════════════════════════════════════════
out = r"e:\BM\project\dev\BlackMist\Tools\GameProxyTool\GameProxyTool使用手册.docx"
doc.save(out)
print(f"已生成：{out}")
